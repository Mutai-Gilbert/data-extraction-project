from docx import Document
import pandas as pd
import os

print("Current working directory:", os.getcwd())

def extract_tables_from_word(file_path):
    """
    Extract data from tables in the Word document and return it as a list of dictionaries.
    """
    table_data = []
    document = Document(file_path)
    
    # Loop through each table in the document to extract data
    for table in document.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            row_data = {headers[i]: cell.text.strip() for i, cell in enumerate(row.cells)}
            table_data.append(row_data)
    
    return table_data

def extract_paragraphs_from_word(file_path):
    """
    Extract paragraphs from the Word document and return them as a list.
    """
    document = Document(file_path)
    paragraph_data = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return paragraph_data

def save_to_excel(tables_data, paragraphs_data, output_file_path):
    """
    Save the extracted data to an Excel file with separate sheets for tables and paragraphs.
    """
    # Ensure the output directory exists
    output_dir = os.path.dirname(output_file_path)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Create a Pandas Excel writer object
    with pd.ExcelWriter(output_file_path) as writer:
        # Save table data to a sheet named 'Tables'
        df_tables = pd.DataFrame(tables_data)
        df_tables.to_excel(writer, sheet_name='Tables', index=False)

        # Save paragraphs to a sheet named 'Paragraphs'
        df_paragraphs = pd.DataFrame(paragraphs_data, columns=['Paragraph'])
        df_paragraphs.to_excel(writer, sheet_name='Paragraphs', index=False)

    print(f"Data successfully saved to {output_file_path}")

if __name__ == "__main__":
    # Define file paths
    input_file = '/Users/mutai/Desktop/data extraction/data/input/sample.docx'
    output_file = '/Users/mutai/Desktop/data extraction/data/output/structured_data.xlsx'
    
    # Extract data from the Word document
    extracted_table_data = extract_tables_from_word(input_file)
    extracted_paragraph_data = extract_paragraphs_from_word(input_file)
    
    # Save the structured data to an Excel file with separate sheets
    save_to_excel(extracted_table_data, extracted_paragraph_data, output_file)
